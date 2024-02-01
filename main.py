"""
HealtheIntent Transformation Exporter to .docx with Syntax Highlighting and Dependency Analysis

Description:
This script processes SQL transformations from HealtheIntent and organises them into folders. Each folder represents a Workflow, containing a Word document for each dataset. 
These documents detail the transformations, with SQL syntax highlighted, and include version and date modified information. 
Transformations that do not contain SQL are excluded
 e.g. cerner-defined workflows and file uploads.

Its main purpose is to format SQL queries from the platform for sharing and to track dependencies between datasets.

Functionality:
    - Converts SQL transformations to .docx files with syntax highlighting.
    - Identifies dependencies between datasets, exporting them to Excel and JSON, and includes recursive dependency tracing.

Usage:
    1. Run 'HealtheIntent Transformation Exporter - All Workflows' from the 'Python Utilities' collection from HealtheIntent's queries page.
    2. Export as CSV, rename to 'transformations.csv'.
    3. Run 'Table Names' from the same collection, export, and rename to 'table_names.csv'.
    4. Place 'transformations.csv' and 'table_names.csv' in the project directory.
    5. Execute the script.
    6. Output will be saved in the 'output' directory.

Notes:
    - Requires libraries: pandas, python-docx, sqlparse, and openpyxl. 
    - Install libraries with 'pip install -r requirements.txt'.

Author: Eddie Davison
Modified: Jan 2024
"""

import pandas as pd
import os
import json
import csv
import logging
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlparse.tokens import Keyword, DML
import time

# Configurable Variables
INPUT_FILENAME = "transformations.csv"
TABLE_NAMES_FILENAME = "table_names.csv"
OUTPUT_DIRNAME = "Output"
PROJECT_DIRECTORY = os.path.dirname(os.path.abspath(__file__))

# Set Paths - leave this as default to use the project directory
input_file_path = os.path.join(PROJECT_DIRECTORY, INPUT_FILENAME)
table_names_file_path = os.path.join(PROJECT_DIRECTORY, TABLE_NAMES_FILENAME)
output_directory_path = os.path.join(PROJECT_DIRECTORY, OUTPUT_DIRNAME)

# Initialize Logger
logging.basicConfig(level=logging.INFO, format='%(message)s')
start_time = time.time()

## Functions

def is_good_sql(transformation_sql):
    """
    Check if the SQL transformation is valid.

    Args:
    transformation_sql (str): The SQL transformation string to be validated.

    Returns:
    bool: True if the transformation SQL contains key elements, False otherwise.
    """
    key_elements = ["SELECT", "FROM"]
    return all(elem in transformation_sql.upper() for elem in key_elements)

def find_all_dependencies(dataset_name, dependency_map, seen=None):
    """
    Recursively find all dependencies for a given dataset.

    Args:
    dataset_name (str): Name of the dataset to find dependencies for.
    dependency_map (dict): Dictionary mapping datasets to their dependencies.
    seen (set, optional): Set of already processed datasets. Defaults to None.

    Returns:
    set: A set of all dependencies for the given dataset.
    """
    if seen is None:
        seen = set()

    if dataset_name not in seen:
        seen.add(dataset_name)
        for dependent in dependency_map.get(dataset_name, []):
            seen.update(find_all_dependencies(dependent, dependency_map, seen))
    return seen

def export_dependencies_to_excel(data, output_path, workflow_data, dependency_type="full"):
    """
    Export dataset dependencies to Excel.

    Args:
    data (dict): Dictionary containing dataset dependencies.
    output_path (str): Path where the Excel file will be saved.
    workflow_data (dict): Dictionary containing workflow and dataset details.
    dependency_type (str): Type of dependencies to export ('full' or 'direct').
    """
    data_list = []
    for dataset, dependencies in data.items():
        workflow_name = workflow_data.get(dataset, {}).get("workflow_name", "UNKNOWN")
        for dependency in dependencies:
            data_list.append((workflow_name, dataset, dependency))

    column_names = ['WORKFLOW_NAME', 'DATA_SET_MNEMONIC', 'DEPENDENCY']
    df = pd.DataFrame(data_list, columns=column_names)
    df.to_excel(output_path, index=False, engine='openpyxl')

def highlight_sql(sql, paragraph):
    """
    Highlight SQL syntax in a Word document paragraph.

    Args:
    sql (str): SQL query string to be highlighted.
    paragraph (docx.text.paragraph.Paragraph): Paragraph object from python-docx containing the transformaqtion SQL.

    Returns:
    None
    """
    
    # SQL keywords
    keywords = ["SELECT", "FROM", "WHERE", "DISTINCT", "AND", "OR", "IN", "NOT IN", 
                "JOIN", "INNER", "LEFT", "RIGHT", "OUTER", "ON", "AS", "LIKE", "GROUP BY", "ORDER BY", 
                "HAVING", "WITH", "EXCLUDE", "UNION", "INTERSECT", "EXCEPT", "CASE", 
                "WHEN", "THEN", "ELSE", "END", "BETWEEN", "OVER", "PARTITION BY", "ROWS", "RANGE"]
    
    # Vertica keywords
    keywords.extend([
        "COPY", "MERGE", "ANALYZE", "COLLECT", "STATISTICS", "PROJECTION", 
        "SEGMENTED", "UNSEGMENTED", "NODES", "REJECTMAX", "ENFORCELENGTH", 
        "TIMEOUT", "LOCAL", "SYSDATE", "SYSTIME", "SYSTIMESTAMP", "CURRENT_DATE", 
        "CURRENT_TIME", "CURRENT_TIMESTAMP", "INTERVAL", "LIMIT", "OFFSET", 
        "OVERLAPS", "USING", "EXCLUSIVE", "SHARED", "EXPLAIN", "PLAN", "PROFILE"
    ])

    # SQL functions
    functions = [
        "COUNT", "SUM", "AVG", "MIN", "MAX", "ROUND", "UPPER", "LOWER", 
        "LENGTH", "LTRIM", "RTRIM", "COALESCE", "CAST", "CONVERT", "CASE",
        "APPROXIMATE_COUNT_DISTINCT", "TO_TIMESTAMP", "DATE_TRUNC", "CASEWHEN",
        "CASE WHEN", "DECODE", "NVL", "NULLIF", "EXTRACT", "POSITION", "SUBSTRING", 
        "WHEN", "CHAR_LENGTH", "OCTET_LENGTH", "TO_CHAR", "TO_NUMBER", "TRIM",
        "LEAD", "LAG", "FIRST_VALUE", "LAST_VALUE", "DENSE_RANK", "NTILE", 
        "PERCENT_RANK", "PERCENTILE_CONT", "PERCENTILE_DISC", "CUME_DIST", 
        "RANK", "ROW_NUMBER", "STDDEV", "STDDEV_POP", "STDDEV_SAMP", 
        "VARIANCE", "VAR_POP", "VAR_SAMP"
    ]

    # Vertica functions
    functions.extend([
        "APPROXIMATE_COUNT_DISTINCT", "APPROXIMATE_MEDIAN", "APPROXIMATE_PERCENTILE", 
        "AUTO_INCREMENT", "BIT_COUNT", "BTRIM", "CURRENT_DATABASE", "CURRENT_SCHEMA", 
        "CURRENT_USER", "DECODE", "ENCODE", "HEX", "INET_ATON", "INET_NTOA", 
        "INITCAP", "ISNULL", "LPAD", "RPAD", "MD5", "RANDOM", "REGEXP_INSTR", 
        "REGEXP_REPLACE", "REGEXP_SUBSTR", "REPLACE", "SPLIT_PART", "TO_DATE", 
        "TO_TIMESTAMP_TZ", "TRANSLATE", "TRUNC", "BINARY", "BOOLEAN", "CHAR", 
        "VARCHAR", "DATE", "TIMESTAMP", "TIMESTAMPTZ", "TIMESTAMP_LTZ", "TIME", 
        "TIME_TZ", "TIME_LTZ", "INTERVAL_YEAR", "INTERVAL_MONTH", "INTERVAL_DAY", 
        "INTERVAL_HOUR", "INTERVAL_MINUTE", "INTERVAL_SECOND", "FLOAT", "REAL", 
        "NUMERIC", "INT", "INTEGER", "BIGINT", "SMALLINT", "TINYINT", "BYTEINT", 
        "HLL_AGGREGATE", "HLL_COMBINE", "HLL_ESTIMATE", "HLL_SYNTHESIZE", "HLL_UNION_AGG",
        "RANK", "DENSE_RANK", "PERCENT_RANK", "CUME_DIST", "NTILE", "ROW_NUMBER",
        "LAG", "LEAD", "FIRST_VALUE", "LAST_VALUE", "LISTAGG", "STRING_AGG"
    ])

    # SQL operators
    operators = ["=", "<", ">", "<=", ">=", "<>", "!=", "||", "+", "-", "*", "/", "%"]

    # Split the SQL into tokens
    tokens = sql.split(" ")
    in_comment = False

    # For each token, add it to the paragraph with the appropriate formatting
    for i, token in enumerate(tokens):
        run = paragraph.add_run(token + " ")
        run.font.name = 'Arial'
        run.font.size = Pt(10)

        # Check if the current token is a comment
        if token.startswith("--") or token.startswith("/*"):
            in_comment = True

        if in_comment:
            run.font.color.rgb = RGBColor(128, 128, 128)  # Grey for comments
            if '\n' in token or token.endswith("*/"):
                in_comment = False
            continue

        # Check if the current token and the next token form a multi-word keyword
        if i < len(tokens) - 1:
            combined_token = f"{token} {tokens[i + 1]}"
            if combined_token.upper() in keywords:
                run.font.color.rgb = RGBColor(128, 0, 128)  # Purple for keywords
                continue

        # Apply formatting based on the token type
        if token.upper() in keywords:
            run.font.color.rgb = RGBColor(128, 0, 128)  # Purple for keywords
        elif token.upper() in functions:
            run.font.color.rgb = RGBColor(0, 128, 0)    # Green for functions
        elif any(op in token for op in operators):
            run.font.color.rgb = RGBColor(255, 0, 0)    # Red for operators

def clean_dependency(value):
    """
    Check if a dependency value is clean (non-empty and not just whitespace).

    Args:
    value (str): The dependency string to be checked.

    Returns:
    bool: True if the value is clean, False otherwise.
    """
    return value.strip() not in ["", "\t", "\t\t\t", "CATEGORY"]

## Main Process Flow
"""
1. Create Output Directory if it doesn't exist.
2. Create a dictionary to store transformations for each workflow and dataset in the CSV file.
3. Extract columns from CSV and create folders for each workflow name.
4. Store transformation SQL in the workflow_transformations dictionary.
5. Extract unique data_set_mnemonics.
6. Read the additional table names from table_names.csv.
7. Combine the unique datasets with the additional table names.
8. Find dependencies for each dataset and store in a dictionary.
9. Save dependencies to JSON.
10. Trace all dependencies recursively.
11. Save full dependencies to JSON.
12. Prepare workflow_data dictionary.
13. Export direct dependencies to Excel.
14. Export full dependencies to Excel.
15. Create a Word document for each dataset.
16. Save the Word document to the appropriate workflow directory.
"""

# Create output directory if it doesn't exist
if not os.path.exists(output_directory_path):
    os.makedirs(output_directory_path)

# Create a dictionary to store transformations for each workflow and dataset in the CSV file
workflow_transformations = {}

# Extract workflow name, dataset name, dataset version, date modified, and transformation SQL
with open(input_file_path, 'r', encoding='utf-8') as csvfile:
    csvreader = csv.reader(csvfile)
    headers = next(csvreader)
    for row in csvreader:
        workflow_name, dataset_name, dataset_version, date_modified, transformation_sql = row
        if is_good_sql(transformation_sql):
            # Define the directory path for the current workflow
            workflow_dir = os.path.join(output_directory_path, workflow_name)
            
            # Create the workflow directory if it doesn't exist
            if not os.path.exists(workflow_dir):
                os.makedirs(workflow_dir)
            
            # Store transformation SQL in the workflow_transformations dictionary
            if workflow_name not in workflow_transformations:
                workflow_transformations[workflow_name] = {}
            if dataset_name not in workflow_transformations[workflow_name]:
                workflow_transformations[workflow_name][dataset_name] = {'version': dataset_version, 'date_modified': date_modified, 'sqls': []}
                
            workflow_transformations[workflow_name][dataset_name]['sqls'].append(transformation_sql)

# Extract unique data_set_mnemonics
unique_data_sets = set()
with open(input_file_path, 'r', encoding='utf-8') as csvfile:
    csvreader = csv.reader(csvfile)
    headers = next(csvreader)
    for row in csvreader:
        _, dataset_name, _, _, _ = row
        unique_data_sets.add(dataset_name.upper())  # Storing in upper case

# Read the additional table names from table_names.csv
additional_table_names = set()
with open(table_names_file_path, 'r', encoding='utf-8') as csvfile:
    csvreader = csv.reader(csvfile)
    headers = next(csvreader)
    for row in csvreader:
        additional_table_names.add(row[0].upper())  # Storing in upper case

# Combine the unique datasets with the additional table names
unique_data_sets = unique_data_sets.union(additional_table_names)

# Find dependencies for each dataset and store in a dictionary
dataset_dependencies = {}

# For each dataset, trace all dependencies
for workflow, datasets in workflow_transformations.items():
    for dataset, details in datasets.items():
        dependencies = set()
        for transformation in details['sqls']:
            upper_transformation = transformation.upper()
            for potential_dependency in unique_data_sets:
                # Check if our potential_dependency appears in the transformation
                if potential_dependency in upper_transformation:
                    if clean_dependency(potential_dependency):
                        dependencies.add(potential_dependency)
        dataset_dependencies[dataset] = list(dependencies)

logging.info(f"Identified dependencies for {len(dataset_dependencies)} datasets.")

# Save dependencies to JSON
dependencies_filepath = os.path.join(output_directory_path, "direct_dependencies.json")
with open(dependencies_filepath, 'w') as json_file:
    json.dump(dataset_dependencies, json_file, indent=4)

# Trace all dependencies recursively
full_dependencies = {}
for dataset in dataset_dependencies.keys():
    full_dependencies[dataset] = list(find_all_dependencies(dataset, dataset_dependencies))
    full_dependencies[dataset].remove(dataset)  # Remove self from dependencies

# Save full dependencies to JSON
full_dependencies_filepath = os.path.join(output_directory_path, "full_dataset_dependencies.json")
with open(full_dependencies_filepath, 'w') as json_file:
    json.dump(full_dependencies, json_file, indent=4)

# Prepare workflow_data dictionary
workflow_data = {}
for workflow, datasets in workflow_transformations.items():
    for dataset in datasets:
        workflow_data[dataset] = {"workflow_name": workflow}

# Export direct dependencies to Excel
direct_dependencies_path = os.path.join(output_directory_path, "direct_dependencies.xlsx")
export_dependencies_to_excel(dataset_dependencies, direct_dependencies_path, workflow_data, dependency_type="direct")
logging.info(f"Exported direct dependencies to {direct_dependencies_path}")

# Export full dependencies to Excel
full_dependencies_path = os.path.join(output_directory_path, "full_dependencies.xlsx")
export_dependencies_to_excel(full_dependencies, full_dependencies_path, workflow_data, dependency_type="full")
logging.info(f"Exported full dependencies to {full_dependencies_path}")

# Initialize counter
word_file_count = 0

# Create a Word document for each dataset
for workflow, datasets in workflow_transformations.items():
    logging.info(f"Processing workflow: {workflow}")
    for dataset, details in datasets.items():
        doc = Document()
        # Add title
        title = doc.add_heading(dataset, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(20)
            run.font.bold = True

         # Add version number
        version_paragraph = doc.add_paragraph(f"Version: {details['version']}")
        for run in version_paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True

        # Add date modified
        date_modified_paragraph = doc.add_paragraph(f"Date Modified: {details['date_modified']}")
        for run in date_modified_paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
        
        # Add each transformation
        for i, transformation in enumerate(details['sqls']):
            heading = doc.add_heading(f'Transformation {i+1}', level=1)
            paragraph = doc.add_paragraph()
            # Highlight the SQL syntax
            highlight_sql(transformation, paragraph)

        # Add list of Direct Dependencies
        direct_deps = dataset_dependencies.get(dataset, [])
        direct_deps_paragraph = doc.add_paragraph()
        direct_deps_paragraph.add_run(f"Direct Dependencies: {len(direct_deps)}\n").bold = True
        for dep in direct_deps:
            direct_deps_paragraph.add_run(dep + "\n")

        # Add list of Full Dependencies
        full_deps = full_dependencies.get(dataset, [])
        full_deps_paragraph = doc.add_paragraph()
        full_deps_paragraph.add_run(f"Full Dependencies: {len(full_deps)}\n").bold = True
        for dep in full_deps:
            full_deps_paragraph.add_run(dep + "\n")

        output_file = os.path.join(output_directory_path, workflow, f"{dataset}.docx")
        doc.save(output_file)
        # Increment the counter
        word_file_count += 1
        logging.info(f"Saved document for dataset {dataset} at {output_file}")

end_time = time.time()
elapsed_time = end_time - start_time
minutes = int(elapsed_time // 60)
seconds = int(elapsed_time % 60)

if minutes == 1:
    logging.info(f"Done! Created {word_file_count} word documents, 2 JSON files and 2 Excel Workbooks in 1 minute and {seconds} seconds")
elif minutes > 1:
    logging.info(f"Done! Created {word_file_count} word documents, 2 JSON files and 2 Excel Workbooks in {minutes} minutes and {seconds} seconds")
else:
    logging.info(f"Done! Created {word_file_count} word documents, 2 JSON files and 2 Excel Workbooks in {seconds} seconds")




